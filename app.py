import os
import re
import streamlit as st
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.shared import qn

try:
    from openai import OpenAI
except Exception:
    OpenAI = None


# =========================
# CONFIG
# =========================

st.set_page_config(page_title="Rahhal CREW", page_icon="⚙️", layout="centered")
DEFAULT_MODEL = "gpt-4.1-mini"


# =========================
# LOAD SYSTEM PROMPT
# =========================

def load_prompt():
    try:
        with open("rahhal_prompt.txt", "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return "Rahhal CREW system active."


RAHHAL_SYSTEM = load_prompt()


# =========================
# OPENAI CLIENT
# =========================

def get_client():
    if OpenAI is None:
        st.sidebar.error("OpenAI library missing.")
        return None

    key = os.getenv("OPENAI_API_KEY")
    if not key:
        st.sidebar.error("OPENAI_API_KEY not configured in environment.")
        return None

    return OpenAI(api_key=key)


# =========================
# DOCX EXPORT (ROBUST TABLES + NO CACHE)
# =========================

def _set_doc_defaults(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    except Exception:
        pass

def _set_paragraph_15_spacing(p):
    p.paragraph_format.line_spacing = 1.5

def _is_md_separator_row(line: str) -> bool:
    s = line.strip().replace(" ", "")
    if "|" not in s:
        return False
    s = s.strip("|")
    parts = s.split("|")
    if len(parts) == 0:
        return False
    for part in parts:
        if part == "":
            continue
        part = part.replace(":", "")
        if not (set(part) <= {"-"} and len(part) >= 3):
            return False
    return True

def _split_md_row(line: str):
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]

def _add_word_table(doc: Document, header, body):
    rows = 1 + len(body)
    cols = len(header)

    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    for j, h in enumerate(header):
        table.cell(0, j).text = h

    for i, row in enumerate(body, start=1):
        for j, val in enumerate(row):
            table.cell(i, j).text = val

    for r in table.rows:
        for cell in r.cells:
            for p in cell.paragraphs:
                _set_paragraph_15_spacing(p)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)

    doc.add_paragraph("")

def export_docx(messages):
    doc = Document()
    _set_doc_defaults(doc)

    doc.add_heading("Rahhal CREW Output", level=1)
    meta = doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    _set_paragraph_15_spacing(meta)

    for m in messages:
        if m.get("role") == "system":
            continue

        role = (m.get("role") or "message").capitalize()
        doc.add_heading(role, level=2)

        content = (m.get("content") or "").replace("\r\n", "\n")
        lines = content.split("\n")

        i = 0
        buffer_paragraph = []

        def flush_paragraph_buffer():
            nonlocal buffer_paragraph
            text = "\n".join([x.rstrip() for x in buffer_paragraph]).strip()
            if text:
                p = doc.add_paragraph(text)
                _set_paragraph_15_spacing(p)
            buffer_paragraph = []

        while i < len(lines):
            line = lines[i]

            # Markdown table start condition:
            # header row contains pipes
            # next row is separator row like | --- | --- |
            if ("|" in (line or "")) and (i + 1 < len(lines)) and _is_md_separator_row(lines[i + 1]):
                flush_paragraph_buffer()

                header = _split_md_row(line)
                i += 2  # skip separator row

                body = []
                while i < len(lines):
                    candidate = lines[i]
                    if candidate.strip() == "":
                        break
                    if _is_md_separator_row(candidate):
                        break
                    if "|" not in candidate:
                        break
                    body.append(_split_md_row(candidate))
                    i += 1

                cols = len(header)
                body = [r + [""] * (cols - len(r)) if len(r) < cols else r[:cols] for r in body]
                _add_word_table(doc, header, body)
                continue

            buffer_paragraph.append(line)
            i += 1

        flush_paragraph_buffer()

    # Unique filename each time to avoid browser caching
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    path = f"rahhal_output_{ts}.docx"
    doc.save(path)
    return path


# =========================
# SESSION INIT
# =========================

if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "system", "content": RAHHAL_SYSTEM},
        {
            "role": "assistant",
            "content": (
                "Rahhal CREW activated.\n\n"
                "Structured crisis readiness exercise design system.\n\n"
                "Select exercise format:\n"
                "Discussion Based Tabletop or Functional Exercise?"
            ),
        },
    ]

if "pending_input" not in st.session_state:
    st.session_state.pending_input = None


# =========================
# HEADER
# =========================

st.markdown("## Rahhal CREW")
st.markdown("Structured Crisis Readiness Exercise Design System")
st.divider()


# =========================
# SIDEBAR
# =========================

with st.sidebar:
    st.header("Control Panel")

    temperature = st.slider("Creativity", 0.0, 1.0, 0.1, 0.05)

    st.divider()

    if st.button("Reset session"):
        del st.session_state.messages
        st.session_state.pending_input = None
        st.rerun()

    st.divider()

    if st.button("Generate full package"):
        st.session_state.pending_input = "FINAL PACKAGE"
        st.rerun()

    st.divider()

    if st.button("Prepare DOCX export"):
        msgs = [m for m in st.session_state.messages if m.get("role") != "system"]
        if msgs:
            file_path = export_docx(msgs)
            with open(file_path, "rb") as f:
                st.download_button(
                    "Download Word file",
                    f,
                    file_name=file_path,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        else:
            st.warning("No content to export.")


# =========================
# TABS
# =========================

tab_chat, tab_help = st.tabs(["Chat", "How it works"])

with tab_help:
    st.markdown(
        """
• Open the Chat tab  
• Select exercise format  
• Respond step by step  
• Click Generate full package when ready  
• Export Word document from sidebar  
"""
    )


# =========================
# CHAT TAB
# =========================

with tab_chat:
    for m in st.session_state.messages:
        if m.get("role") == "assistant":
            with st.chat_message("assistant", avatar="⚙️"):
                st.write(m.get("content", ""))
        elif m.get("role") == "user":
            with st.chat_message("user", avatar="👤"):
                st.write(m.get("content", ""))

    user_text = st.chat_input("Write your response")
    if user_text:
        st.session_state.pending_input = user_text
        st.rerun()


# =========================
# PROCESS MODEL TURN SAFELY
# =========================

if st.session_state.pending_input:
    user_text = st.session_state.pending_input
    st.session_state.pending_input = None

    st.session_state.messages.append({"role": "user", "content": user_text})

    client = get_client()
    if client:
        try:
            response = client.chat.completions.create(
                model=DEFAULT_MODEL,
                temperature=temperature,
                messages=st.session_state.messages,
            )
            reply = response.choices[0].message.content
            st.session_state.messages.append({"role": "assistant", "content": reply})
        except Exception as e:
            st.sidebar.error(f"API error: {e}")

    st.rerun()
